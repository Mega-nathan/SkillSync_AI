from typing import Union

def extract_text_from_docx(path_or_bytes: Union[str, bytes]) -> str:
	try:
		import docx
	except Exception:
		return ""

	try:
		if isinstance(path_or_bytes, (bytes, bytearray)):
			from io import BytesIO

			doc = docx.Document(BytesIO(path_or_bytes))
		else:
			doc = docx.Document(path_or_bytes)
		paragraphs = [p.text for p in doc.paragraphs]
		return "\n".join(paragraphs)
	except Exception:
		return ""
